from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# Create a new Document
doc = Document()

# Define font name for consistency
font_name = 'Times New Roman'

# Function to add a new page for each experiment
def add_page_break():
    doc.add_page_break()

# Function to add heading with specific style and size
def add_heading(text, size, is_bold=False, is_underlined=False, alignment='left'):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.bold = is_bold
    run.underline = is_underlined
    run.font.name = font_name
    if alignment == 'center':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Function to add normal text
def add_text(text, size):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name

# Add content for each program
def add_program(program_number, aim, code, output=None, yacc_code=None):
    add_heading(f"Program - {program_number}", 14, is_bold=True, is_underlined=True, alignment='center')
    add_heading("Aim", 12, is_bold=True, is_underlined=True)
    add_text(aim, 12)
    add_heading("Code", 12, is_bold=True, is_underlined=True)
    add_text(code, 12)
    if yacc_code:
        add_heading("YACC Code", 12, is_bold=True, is_underlined=True)
        add_text(yacc_code, 12)
    if output:
        add_heading("Output", 12, is_bold=True, is_underlined=True)
        add_text(output, 12)

# Sample Programs

# Program 9 - Left Factoring
program9_aim = "Write a program to perform Left Factoring on a Grammar."
program9_code = """
#include <stdio.h>
#include <string.h>

void leftFactoring(char* nonTerminal, char* productions) {
    char commonPrefix[100], remaining1[100], remaining2[100];
    char* prod1 = strtok(productions, "|");
    char* prod2 = strtok(NULL, "|");

    int i = 0;
    while (prod1[i] == prod2[i] && prod1[i] != '\\0' && prod2[i] != '\\0') {
        commonPrefix[i] = prod1[i];
        i++;
    }
    commonPrefix[i] = '\\0';

    if (strlen(commonPrefix) > 0) {
        printf("%s -> %s%s'\\n", nonTerminal, commonPrefix, nonTerminal);
        printf("%s' -> %s | %s\\n", nonTerminal, prod1 + i, prod2 + i);
    } else {
        printf("No left factoring needed.\\n");
    }
}
"""
program9_output = """
Example Input: A -> ab|ac
Example Output:
A -> aA'
A' -> b | c
"""

# Program 10 - Stack Operations
program10_aim = "Write a program to show all the operations of a stack."
program10_code = """
#include <stdio.h>
#define MAX 5

int stack[MAX], top = -1;

void push(int x) {
    if (top < MAX - 1) {
        stack[++top] = x;
        printf("Pushed %d onto stack\\n", x);
    } else {
        printf("Stack Overflow\\n");
    }
}

void pop() {
    if (top >= 0) {
        printf("Popped %d from stack\\n", stack[top--]);
    } else {
        printf("Stack Underflow\\n");
    }
}

void display() {
    if (top >= 0) {
        printf("Stack: ");
        for (int i = 0; i <= top; i++) {
            printf("%d ", stack[i]);
        }
        printf("\\n");
    } else {
        printf("Stack is empty\\n");
    }
}
"""
program10_output = """
Operations:
push(10), push(20), pop(), display()

Expected Output:
Pushed 10 onto stack
Pushed 20 onto stack
Popped 20 from stack
Stack: 10
"""

# Program 11 - Leading of Non-Terminals
program11_aim = "Write a program to find out the leading of the non-terminals in a grammar."
program11_code = """
#include <stdio.h>
#include <string.h>

void findLeading(char nonTerminal, char productions[10][10], int count) {
    printf("Leading(%c) = {", nonTerminal);
    for (int i = 0; i < count; i++) {
        if (productions[i][0] == nonTerminal) {
            if (productions[i][2] >= 'a' && productions[i][2] <= 'z') {
                printf(" %c", productions[i][2]);
            }
        }
    }
    printf(" }\\n");
}
"""

# Program 12 - Shift Reduce Parsing
program12_aim = "Write a program to Implement Shift Reduce parsing for a String."
program12_code = """
#include <stdio.h>
#include <string.h>

char stack[50];
int top = -1;

void push(char symbol) {
    stack[++top] = symbol;
}

void pop() {
    top--;
}

void display() {
    for (int i = 0; i <= top; i++) {
        printf("%c", stack[i]);
    }
}

void check() {
    if (stack[top] == 'b' && stack[top-1] == 'S' && stack[top-2] == 'a') {
        pop(); pop(); pop();
        push('S');
        printf(" Reduced by: S -> aSb\\n");
        display();
    }
}

int main() {
    char input[20];
    printf("Enter the input string: ");
    scanf("%s", input);
    for (int i = 0; i < strlen(input); i++) {
        push(input[i]);
        printf(" Shift\\n");
        check();
    }
    return 0;
}
"""

# Program 13 - FIRST of Non-Terminals
program13_aim = "Write a program to find out the FIRST of the Non-terminals in a grammar."
program13_code = """
#include <stdio.h>
#include <string.h>

void findFirst(char nonTerminal, char productions[10][10], int count) {
    printf("First(%c) = {", nonTerminal);
    for (int i = 0; i < count; i++) {
        if (productions[i][0] == nonTerminal) {
            printf(" %c", productions[i][2]);
        }
    }
    printf(" }\\n");
}
"""

# Program 14 - Operator Precedence Grammar
program14_aim = "Write a program to check whether a grammar is operator precedence."
program14_code = """
#include <stdio.h>
#include <string.h>
#include <stdbool.h>

bool isOperatorPrecedence(char* grammar) {
    for (int i = 0; grammar[i] != '\\0'; i++) {
        if ((grammar[i] == '+' || grammar[i] == '*') && grammar[i + 1] == grammar[i]) {
            return false;
        }
    }
    return true;
}

int main() {
    char grammar[100];
    printf("Enter the grammar rules: ");
    scanf("%s", grammar);
    if (isOperatorPrecedence(grammar)) {
        printf("The grammar is operator precedence.\\n");
    } else {
        printf("The grammar is not operator precedence.\\n");
    }
}
"""
program14_output = """
Example Input: E -> E+E|E*E|id
Example Output: The grammar is operator precedence.
"""

# Adding all programs to the document, with a page break after each
add_program(9, program9_aim, program9_code, output=program9_output)
add_page_break()

add_program(10, program10_aim, program10_code, output=program10_output)
add_page_break()

add_program(11, program11_aim, program11_code)
add_page_break()

add_program(12, program12_aim, program12_code)
add_page_break()

add_program(13, program13_aim, program13_code)
add_page_break()

add_program(14, program14_aim, program14_code, output=program14_output)

# Save the document
doc_path = "programs_9_to_14.docx"
doc.save(doc_path)

doc_path
